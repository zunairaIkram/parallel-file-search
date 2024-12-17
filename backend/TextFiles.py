from flask import Flask, request, jsonify
from concurrent.futures import ProcessPoolExecutor, as_completed
import os, re
import io
from docx import Document
from PyPDF2 import PdfReader
import fitz



def process_chunk(chunk, start_line, pattern, fileName):
    matches = []
    regex = re.compile(pattern, re.IGNORECASE)
    # print(f"Process {os.getpid()} starting to process file '{fileName}' from line {start_line}")
    
    for i, line in enumerate(chunk):
        if regex.search(line):
            matches.append({
                "lineNumber": start_line + i,
                "line": line.strip(),
                "processId": os.getpid(),
                "fileName": fileName
            })
            # print(f"Process {os.getpid()} found pattern in file '{fileName}', line {start_line + i}")
    
    # print(f"Process {os.getpid()} finished processing file '{fileName}' up to line {start_line + len(chunk)}")
    return matches



def divide_into_chunks(lines, chunk_size):
    for i in range(0, len(lines), chunk_size):
        yield lines[i:i + chunk_size]



def parallel_search_in_file_content(content, pattern, fileName):
    try:
        # print(f"Process {os.getpid()} starting to process file '{fileName}'")

        lines = []
        file_ext = os.path.splitext(fileName)[1].lower() 
        
        if file_ext == '.txt':
            lines = content.decode('utf-8').splitlines()
            # lines = extraction_from_text(content)
        elif file_ext == '.docx':
            lines = extraction_from_docx(content)
        elif file_ext == '.pdf':
            lines = extraction_from_pdfs(content)
        else:
            return jsonify({"error": f"Unsupported file type: {file_ext}"}), 400
        
        total_lines = len(lines)
        total_cores = os.cpu_count()
        chunk_size = max(1, total_lines // total_cores)

        if total_lines % total_cores != 0:
            chunk_size += 1

        results = []
        with ProcessPoolExecutor() as executor:
            futures = []
            start_line = 1

            for chunk in divide_into_chunks(lines, chunk_size):
                futures.append(executor.submit(process_chunk, chunk, start_line, pattern, fileName))
                start_line += len(chunk)

            for future in as_completed(futures):
                matches = future.result()
                if matches:
                    results.extend(matches)
        
                    
        # print(f"Process {os.getpid()} ending to process file '{fileName}'")
        
        return results

    except Exception as e:
        print(f"An error occurred: {e}")
        return []



def parallel_search_in_multiple_files(file_contents, pattern):
    try:
        results = []
        with ProcessPoolExecutor() as executor:
            futures = []
            
            for file_content in file_contents:
                futures.append(
                    executor.submit(
                        parallel_search_in_file_content,
                        file_content['content'],
                        pattern,
                        file_content['filename']
                    )
                )

            for future in as_completed(futures):
                matches = future.result()
                if matches:
                    results.append({
                        "fileName": matches[0]["fileName"],
                        "matches": matches
                    })
        
        return results

    except Exception as e:
        print(f"An error occurred in multiple file search: {e}")
        return []



def text_files_searching():
    try:
        pattern = request.form.get("pattern")
        uploaded_files = request.files.getlist("files")

        if not pattern:
            return jsonify({"error": "No search pattern provided"}), 400
        
        if not uploaded_files:
            return jsonify({"error": "No files uploaded"}), 400

        file_contents = []

        for uploaded_file in uploaded_files:
            try:
                filename = uploaded_file.filename
                content = uploaded_file.read()

                file_contents.append({"filename": filename, "content": content})

            except UnicodeDecodeError:
                return jsonify({"error": f"Could not decode file: {filename}"}), 400
            except Exception as e:
                return jsonify({"error": f"Error processing file {filename}: {str(e)}"}), 400

        results = parallel_search_in_multiple_files(file_contents, pattern)
        return jsonify(results)

    except Exception as e:
        print(f"An error occurred: {e}")
        return jsonify({"error": "An internal error occurred"}), 500









# def extraction_from_text(content):
#     # Decode the content (assuming it's in bytes)
#     text = content.decode('utf-8')  
#     # Split text into lines
#     all_content = text.splitlines()
#     # Process each line to split by sentence-ending punctuation
#     final_content = []
#     for line in all_content:
#         # Split sentences by periods (with trailing space) and retain the period
#         split_lines = re.split(r'(?<=\.)\s+', line)
#         final_content.extend(split_lines)

#     return final_content



def extraction_from_docx(content):
    all_content = []

    try:
        doc = Document(io.BytesIO(content))

        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text.strip()
            if paragraph_text:
                all_content.append(paragraph_text)

        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text) 
                if row_text:
                    all_content.append("    ".join(row_text))


        for shape in doc.inline_shapes:
            if shape.type == 1: 
                textbox_doc = shape._element
                if textbox_doc is not None:
                    for paragraph in textbox_doc.xpath(".//w:p"):
                        texts = [node.text for node in paragraph.findall(".//w:t") if node.text]
                        if texts:
                            shape_text = "".join(texts).strip()
                            all_content.append(shape_text)

    except Exception as e:
        all_content.append(f"[Error processing .docx file: {str(e)}]")

    final_content = []
    for line in all_content:
        split_lines = re.split(r'(?<=\.)\s+', line)
        final_content.extend(split_lines)

    return final_content



def extraction_from_pdfs(content):
    combined_content = []

    try:
        pdf_document = fitz.open("pdf", content)

        for page in pdf_document:
            blocks = page.get_text("dict")["blocks"] 
            
            for block in blocks:
                if "lines" in block and block.get("type", -1) == 0: 
                    for line in block["lines"]:
                        line_text = ""
                        for span in line["spans"]:
                            span_text = span["text"].strip()
                            if span_text: 
                                if line_text:
                                    line_text += " " 
                                line_text += span_text
                        if line_text:  
                            combined_content.append(line_text)

        full_text = " ".join(combined_content) 

        final_content = re.split(r'(?<=\.)\s+', full_text)

    except Exception as e:
        raise ValueError(f"Error extracting text from PDF: {str(e)}")

    return final_content


